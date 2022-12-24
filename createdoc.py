"""
    TODO:
    - Asynchronous - do multiple pages at a time
    - Format the doc nicer
    - Explicitly tell GPT to give title to each section/page given, and put this in docx in bold and the paragaroh subheading
"""
import argparse
import re
import openai
from pyChatGPT import ChatGPT
import tomli
from docx import Document
from docx.shared import Inches
import requests
import io



class BookCreator:

    def __init__(self):
        with open("config.toml", mode="rb") as fp:
            self.config = tomli.load(fp)

        self.chatapi = ChatGPT(session_token=self.config['chatGPT']['session_token'])
        openai.api_key = self.config['openapi']['apikey']



    def document_create(self, inputfile_name: str, booktitle: str):
        document = self.document_setup(booktitle)

        with open(inputfile_name, 'r', encoding="utf8") as inputfile:
            inputtext = inputfile.read()


        regex_pattern_newline = r'\n([a-zA-Z0-9])'
        inputtext = re.sub(regex_pattern_newline, r" \1", inputtext)

        regex_pattern_page = r'Page \| \d+ Harry Potter and the Philosophers Stone - J\.K\. Rowling\n\n\n'
        splittext = re.split(regex_pattern_page, inputtext)


        try:
            for page in splittext:
                p = document.add_paragraph()
                r = p.add_run()

                resp = self.chatapi\
                    .send_message(f'{self.config["prompts"]["rewritting_style"]}, rewrite the following:\n{page}')
                r.add_text(f'{resp["message"]}\n')

                r = p.add_run()
                resp = self.chatapi\
                    .send_message('Provide a description of what this scene looks like, so I can get an AI image creator to draw it')
                r.add_picture(
                    self.create_image(resp['message']), width=Inches(5))

                self.chatapi.reset_conversation()  # reset the conversation

            print('Complete')

        except Exception as e:
            print(f'Incomplete due to error:\t {e}')

        finally:
            booktitle_path = booktitle\
                .replace(' ', '')\
                .replace('-', '_')

            document.save(f'output_files/{booktitle_path}.docx')


    def document_setup(self, booktitle: str):
        document = Document()
        document.add_heading(booktitle, 0)

        return document



    def create_image(self, description: str):
        resultJSON = openai.Image.create(
            prompt=description,
            n=self.config['openapi']['image_num'],
            size=self.config['openapi']['image_size']
        )

        response = requests.get(resultJSON['data'][0]['url'], stream=True)
        image = io.BytesIO(response.content)
        return image




# MAIN FUNC
def main(args):
    creator = BookCreator()
    creator.document_create(args.input_file, args.book_name)


# ACCEPTING ARGS
if __name__ == "__main__":
    print("HEHEHEHEHE")
    parser = argparse.ArgumentParser()

    parser.add_argument("-i", "--input_file", required=True,
                        help="Name of input file")

    parser.add_argument("-n", "--book_name", required=True,
                        help="Name of output book")

    args = parser.parse_args()
    main(args)


